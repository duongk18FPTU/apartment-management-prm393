from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

import generate_srs_template_aligned as base


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SRS_SECTION_II_TEMPLATE_ALIGNED.md"
MARKDOWN_TARGET = ROOT / "docs" / "SRS_SECTION_II_FINAL_PASTE_MOCKUPS_V2.md"
DOCX_TARGET = ROOT / "docs" / "SRS_SECTION_II_FINAL_PASTE_MOCKUPS_V2.docx"


UC_TO_GROUP = {
    "UC-01": "E.1",
    "UC-02": "E.2, E.3, and E.4",
    "UC-03": "E.1",
    "UC-04": "E.1",
    "UC-05": "E.2",
    "UC-06": "E.2",
    "UC-07": "E.10",
    "UC-08": "E.10",
    "UC-09": "E.2",
    "UC-10": "E.5",
    "UC-11": "E.5",
    "UC-12": "E.5",
    "UC-13": "E.7",
    "UC-14": "E.7",
    "UC-15": "E.7",
    "UC-16": "E.6",
    "UC-17": "E.6",
    "UC-18": "E.6",
    "UC-19": "E.6",
    "UC-20": "E.6",
    "UC-21": "E.6",
    "UC-22": "E.6",
    "UC-23": "E.9",
    "UC-24": "E.9",
    "UC-25": "E.8",
    "UC-26": "E.8",
    "UC-27": "E.8",
}


GROUPS = [
    {
        "id": "E.1",
        "name": "Authentication and Profile",
        "use_cases": "UC-01, UC-03, UC-04",
        "frames": [
            "Login",
            "Change Password",
            "User Profile",
        ],
        "content": [
            "Application logo or building illustration.",
            "Email and masked Password fields.",
            "Password visibility controls.",
            "Login button, loading state, inline validation, and authentication-error area.",
            "Current Password, New Password, and Confirm New Password on Change Password.",
            "Profile fields: Email, Full Name, Phone, National ID, and Date of Birth.",
            "Save Profile, theme control, and Logout.",
        ],
        "exclude": [
            "Username-only login.",
            "Close button.",
            "Browser URL bar.",
            "LMS labels.",
            "Forgot Password, OTP, or self-service password recovery.",
        ],
    },
    {
        "id": "E.2",
        "name": "Admin Dashboard and Administration",
        "use_cases": "UC-02, UC-05, UC-06, UC-09",
        "frames": [
            "Admin Dashboard",
            "User List",
            "User Create or Edit",
            "Resident List",
            "Resident Form or Profile",
        ],
        "content": [
            "Admin Dashboard app bar with Refresh and Logout.",
            "Counters: Apartments, Residents, Pending Requests, Unpaid Bills, Visitors Inside.",
            "Quick access to User Accounts, Apartments, Requests, Complaints, Visitors, and Announcements.",
            "Bottom navigation: Dashboard, Apartments, Residents, Profile.",
            "User search, role filter, status filter, and user cards.",
            "User form fields for email, identity, role, apartment, and status.",
            "Resident search, status filter, identity information, apartment assignment, and profile details.",
        ],
        "exclude": [
            "BQL Manager, BQL Staff, BQT Head, or BQT Member roles.",
            "Export Report controls.",
            "System Settings, Audit Logs, or Backup and Restore.",
        ],
    },
    {
        "id": "E.3",
        "name": "Staff Home",
        "use_cases": "UC-02 and the entry points for UC-12, UC-15, UC-16, UC-17, UC-20, UC-21, UC-24, UC-26, UC-27",
        "frames": ["Staff Home"],
        "content": [
            "Indexed content area.",
            "Bottom navigation: Requests, Bills, Visitors, Complaints, Profile.",
            "Announcement floating action button.",
            "Logout in Profile.",
        ],
        "exclude": [
            "Apartment or user-account administration tabs.",
            "Payment Gateway or wallet controls.",
        ],
    },
    {
        "id": "E.4",
        "name": "Resident Home",
        "use_cases": "UC-02 and the entry points for UC-10, UC-11, UC-13, UC-14, UC-18, UC-19, UC-22, UC-23, UC-25",
        "frames": ["Resident Home"],
        "content": [
            "Resident greeting or apartment summary.",
            "Quick actions: Bills, Maintenance Requests, Complaints, Announcements.",
            "Visitor Registration card.",
            "Payment History card.",
            "Bottom navigation matching the implemented Resident home.",
            "Profile and Logout access.",
        ],
        "exclude": [
            "Admin or Staff management actions.",
            "Wallet balance or automatic deduction controls.",
        ],
    },
    {
        "id": "E.5",
        "name": "Maintenance Request",
        "use_cases": "UC-10, UC-11, UC-12",
        "frames": [
            "Request Create",
            "My Request List",
            "Request Details",
            "Request Management",
        ],
        "content": [
            "Title, Category selector, Description, and image picker.",
            "Image preview list, remove actions, and maximum-three-image indicator.",
            "Submit Request and inline validation.",
            "Resident request cards and status badges.",
            "Request details, images, processing account, and resolution note.",
            "Admin or Staff status filters, status dropdown, resolution input, and Update Status.",
        ],
        "exclude": [
            "Confirmed status.",
            "A separate Service Order module.",
        ],
    },
    {
        "id": "E.6",
        "name": "Bill and Manual Payment",
        "use_cases": "UC-16, UC-17, UC-18, UC-19, UC-20, UC-21, UC-22",
        "frames": [
            "Staff Bill List and Bill Create",
            "Staff Bill Details and Payment Verification",
            "Resident My Bills and Bill Payment",
            "Resident Payment History",
        ],
        "content": [
            "Bill filters by apartment, month, and status.",
            "Bill Create fields: apartment, type, amount, billing month, and due date.",
            "Bill Details: type, apartment, amount, month, due date, and status.",
            "Record Cash Payment.",
            "Pending bank-transfer information, Approve, Reject, and rejection-reason dialog.",
            "Resident bill cards, amount, due date, and status.",
            "Building bank name, account number, transfer reference, and copy actions.",
            "Transfer-confirmation button and text explaining Staff verification.",
            "Payment History with amount, method, status, and date-time.",
        ],
        "exclude": [
            "VNPay, MoMo, or another Payment Gateway.",
            "In-app wallet.",
            "Automatic monthly deduction.",
            "Automatic refund.",
        ],
    },
    {
        "id": "E.7",
        "name": "Complaint and Feedback",
        "use_cases": "UC-13, UC-14, UC-15",
        "frames": [
            "Complaint Create and My Complaint List",
            "Complaint Details",
            "Complaint Management",
        ],
        "content": [
            "Complaint or feedback content field and Submit Complaint.",
            "Resident complaint list and status badges.",
            "Complaint details, response, responder, and response time.",
            "Admin or Staff status filter and complaint list.",
            "Mark In Review, response input, and Resolve.",
        ],
        "exclude": [
            "A separate generic Feedback system unrelated to complaints.",
        ],
    },
    {
        "id": "E.8",
        "name": "Visitor Management",
        "use_cases": "UC-25, UC-26, UC-27",
        "frames": [
            "Resident Visitor Registration",
            "Admin or Staff Visitor List and Check-In or Check-Out",
        ],
        "content": [
            "Visitor name, phone number, purpose, expected time, and apartment.",
            "Register Visitor.",
            "Visitor search and visitor cards.",
            "Registered, Checked In, and Checked Out status badges.",
            "Check In action for Registered.",
            "Check Out action for Checked In.",
            "Check-in and check-out timestamps.",
        ],
        "exclude": [
            "Face verification.",
            "Vehicle registration.",
        ],
    },
    {
        "id": "E.9",
        "name": "Announcement",
        "use_cases": "UC-23, UC-24",
        "frames": [
            "Announcement List",
            "Announcement Create or Edit",
            "Announcement Details",
        ],
        "content": [
            "Announcement cards with title, type, and date.",
            "Create button visible only to Admin and Staff.",
            "Title, Content, Type, Target Roles, and Save.",
            "Details with title, content, type, creator, and time.",
            "Edit and Delete visible only to Admin and Staff.",
            "Delete-confirmation dialog.",
        ],
        "exclude": [
            "A claim that FCM push delivery is already implemented.",
            "Resident create, edit, or delete controls.",
        ],
    },
    {
        "id": "E.10",
        "name": "Apartment Management",
        "use_cases": "UC-07, UC-08",
        "frames": [
            "Apartment List",
            "Apartment Form",
            "Apartment Details and Resident Picker",
        ],
        "content": [
            "Apartment search, floor filter, and occupancy-status filter.",
            "Apartment number, floor, building, area, type, price, and status.",
            "Apartment owner section and Resident member list.",
            "Assign Resident and Resident Picker dialog.",
            "Edit.",
            "Delete with confirmation.",
        ],
        "exclude": [
            "Contract Management.",
            "Amenity Booking.",
        ],
    },
]

# Final representative mock-up set approved for the course submission. Each item
# below corresponds to one hand-drawn image or image board supplied separately.
GROUPS = [
    {
        "id": "E.1",
        "name": "Login",
        "use_cases": "UC-01",
        "frames": ["Login"],
        "content": [
            "App logo or building illustration.",
            "Title: Welcome Back or the title currently used by the application.",
            "Email field with email icon.",
            "Password field with visibility toggle.",
            "Login button.",
            "Loading state on the Login button.",
            "Area for inline validation or authentication error.",
        ],
        "exclude": [
            "Username field.",
            "Close button.",
            "Browser URL bar.",
            "LMS label.",
        ],
    },
    {
        "id": "E.2",
        "name": "Admin Dashboard",
        "use_cases": "UC-02, UC-05, UC-06, and UC-09",
        "frames": ["Admin Dashboard"],
        "content": [
            "App bar with Admin Dashboard, Refresh, and Logout.",
            "Five counters: Apartments, Residents, Pending Requests, Unpaid Bills, and Visitors Inside.",
            "Quick access: User Accounts, Apartments, Maintenance Requests, Complaints, Visitors, and Announcements.",
            "Bottom navigation: Dashboard, Apartments, Residents, and Profile.",
        ],
        "exclude": [
            "Functions or role labels that do not exist in the current application.",
            "Report export, system settings, audit logs, backup, or restore controls.",
        ],
    },
    {
        "id": "E.3",
        "name": "Staff Home",
        "use_cases": "UC-02 and Staff entry points",
        "frames": ["Staff Home"],
        "content": [
            "Indexed content area.",
            "Bottom navigation: Requests, Bills, Visitors, Complaints, and Profile.",
            "Announcement floating action button.",
        ],
        "exclude": [
            "Admin-only apartment or user-account administration.",
            "Payment Gateway or wallet controls.",
        ],
    },
    {
        "id": "E.4",
        "name": "Resident Home",
        "use_cases": "UC-02 and Resident entry points",
        "frames": ["Resident Home"],
        "content": [
            "Resident greeting or apartment summary.",
            "Quick actions: Bills, Maintenance Requests, Complaints, and Announcements.",
            "Visitor Registration card.",
            "Payment History card.",
            "Bottom navigation matching the implemented Resident Home structure.",
        ],
        "exclude": [
            "Admin or Staff management actions.",
            "Wallet balance or automatic deduction controls.",
        ],
    },
    {
        "id": "E.5",
        "name": "Request Create",
        "use_cases": "UC-10, representative of the Maintenance Request feature",
        "frames": ["Request Create"],
        "content": [
            "Title.",
            "Category selector: Plumbing, Electrical, and General.",
            "Description.",
            "Image picker and image preview list.",
            "Submit button.",
            "Inline validation.",
        ],
        "exclude": [
            "Confirmed status.",
            "A separate Service Order module.",
        ],
    },
    {
        "id": "E.6",
        "name": "Bill Details and Payment Verification",
        "use_cases": "UC-17, UC-19, UC-20, and UC-21",
        "frames": ["Staff version", "Resident version"],
        "content": [
            "Staff version: bill type, apartment, amount, month, due date, and status.",
            "Staff version: pending payment information when available.",
            "Staff version: Record Cash Payment.",
            "Staff version: Approve Bank Transfer.",
            "Staff version: Reject Bank Transfer.",
            "Staff version: reject-reason dialog.",
            "Resident version: bill information.",
            "Resident version: building bank name and account number.",
            "Resident version: transfer reference.",
            "Resident version: copy actions.",
            "Resident version: confirmation button.",
            "Resident version: text explaining that Staff must verify the transfer.",
        ],
        "exclude": [
            "VNPay, MoMo, or another Payment Gateway.",
            "In-app wallet, automatic deduction, or automatic refund.",
        ],
    },
    {
        "id": "E.7",
        "name": "Complaint Management",
        "use_cases": "UC-13, UC-14, and UC-15",
        "frames": ["Complaint Management"],
        "content": [
            "Status filter.",
            "Complaint list.",
            "Complaint content and Resident or apartment information.",
            "Mark In Review action.",
            "Response input.",
            "Resolve action.",
        ],
        "exclude": [
            "A separate feedback system unrelated to complaints.",
        ],
    },
    {
        "id": "E.8",
        "name": "Visitor Management",
        "use_cases": "UC-25, UC-26, and UC-27",
        "frames": ["Resident version", "Staff or Admin version"],
        "content": [
            "Resident version: visitor name.",
            "Resident version: phone number.",
            "Resident version: purpose.",
            "Resident version: expected time.",
            "Resident version: Register button.",
            "Staff or Admin version: search.",
            "Staff or Admin version: visitor list.",
            "Staff or Admin version: status badge.",
            "Staff or Admin version: Check-In action for Registered.",
            "Staff or Admin version: Check-Out action for Checked In.",
            "Staff or Admin version: check-in and check-out timestamps.",
        ],
        "exclude": [
            "Face verification.",
            "Vehicle registration.",
        ],
    },
    {
        "id": "E.9",
        "name": "Announcement",
        "use_cases": "UC-23 and UC-24",
        "frames": ["Announcement List", "Announcement Create or Edit", "Announcement Details"],
        "content": [
            "List: announcement cards with title, type, and date.",
            "List: Create button visible only to Admin and Staff.",
            "Create or Edit: Title.",
            "Create or Edit: Content.",
            "Create or Edit: Type.",
            "Create or Edit: Target Roles.",
            "Create or Edit: Save button.",
            "Details: title, content, type, creator, and time.",
            "Details: Edit and Delete visible only to Admin and Staff.",
        ],
        "exclude": [
            "A claim that push-notification delivery is already implemented.",
            "Resident create, edit, or delete controls.",
        ],
    },
    {
        "id": "E.10",
        "name": "Apartment Details",
        "use_cases": "UC-07 and UC-08",
        "frames": ["Apartment Details"],
        "content": [
            "Apartment number, floor, building, area, type, price, and status.",
            "Owner section.",
            "Resident member list.",
            "Assign Resident.",
            "Edit.",
            "Delete with confirmation.",
        ],
        "exclude": [
            "Contract Management.",
            "Amenity Booking.",
        ],
    },
]


def build_group_section() -> str:
    lines = [
        "## E. Screen Mock-Up Redrawing Guide",
        "",
        "The ten groups below define the representative hand-drawn images. Paste the completed images into the centered insertion frames provided in the related Screen Mock-up subsections of Section 3.",
        "",
    ]
    for group in GROUPS:
        lines.extend(
            [
                f"### {group['id']} {group['name']}",
                "",
                f"**Use cases covered:** {group['use_cases']}",
                "",
                "**Frames to include:**",
                "",
            ]
        )
        lines.extend(f"- {frame}" for frame in group["frames"])
        lines.extend(
            [
                "",
                "**Required content:**",
                "",
            ]
        )
        lines.extend(f"- {item}" for item in group["content"])
        lines.extend(["", "**Do not draw:**", ""])
        lines.extend(f"- {item}" for item in group["exclude"])
        lines.append("")
    return "\n".join(lines).rstrip()


def replace_function_boxes(markdown: str) -> str:
    count = markdown.count("[[FIGURE_MOCKUP|")
    if count != 27:
        raise RuntimeError(f"Expected 27 function mock-up markers, found {count}.")
    return markdown


def replace_appendix_e(markdown: str) -> str:
    start_heading = "## E. Screen Mock-Up Redrawing"
    end_heading = "# Final Consistency Checklist"
    start = markdown.index(start_heading)
    end = markdown.index(end_heading)
    return markdown[:start] + build_group_section() + "\n\n" + markdown[end:]


def validate(markdown: str) -> None:
    if markdown.count("[[FIGURE_MOCKUP|") != 27:
        raise RuntimeError("Expected 27 figure mock-up insertion markers.")
    for uc_number in range(1, 28):
        uc_id = f"UC-{uc_number:02d}"
        if uc_id not in markdown:
            raise RuntimeError(f"Missing {uc_id}.")

    with ZipFile(DOCX_TARGET) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"Corrupt DOCX member: {bad_member}")

    document = Document(DOCX_TARGET)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    if text.count("DÁN ẢNH Ở ĐÂY") != 27:
        raise RuntimeError("Expected 27 DOCX image insertion boxes.")
    for removed_text in (
        "Screen or UI scope:",
        "Drawing requirements:",
        "Mock-up reference:",
    ):
        if removed_text in text:
            raise RuntimeError(f"Removed text remains in DOCX: {removed_text}")
    if sum(
        1
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Alternative flows"
    ) != 27:
        raise RuntimeError("Expected 27 Alternative flows sections.")
    if sum(
        1
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Business Rules"
    ) != 27:
        raise RuntimeError("Expected 27 Business Rules sections.")

    print(f"Markdown: {MARKDOWN_TARGET}")
    print(f"DOCX: {DOCX_TARGET}")
    print(f"DOCX size: {DOCX_TARGET.stat().st_size} bytes")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")
    print("Drawing-guide groups: 10")
    print("Use-case image insertion frames: 27")


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    markdown = replace_function_boxes(markdown)
    markdown = replace_appendix_e(markdown)
    MARKDOWN_TARGET.write_text(markdown, encoding="utf-8")

    original_docx_target = base.DOCX_TARGET
    try:
        base.DOCX_TARGET = DOCX_TARGET
        base.build_docx(markdown)
    finally:
        base.DOCX_TARGET = original_docx_target

    validate(markdown)


if __name__ == "__main__":
    main()
