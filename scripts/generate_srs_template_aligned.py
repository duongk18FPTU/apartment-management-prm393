from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SRS_SECTION_II_REVISED_DRAFT.md"
MARKDOWN_TARGET = ROOT / "docs" / "SRS_SECTION_II_TEMPLATE_ALIGNED.md"
DOCX_TARGET = ROOT / "docs" / "SRS_SECTION_II_TEMPLATE_ALIGNED.docx"


def split_markdown_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [part.strip().replace(r"\|", "|") for part in value.split("|")]


def read_table_after(lines: list[str], header_prefix: str) -> list[list[str]]:
    header_index = next(
        index for index, line in enumerate(lines) if line.startswith(header_prefix)
    )
    rows: list[list[str]] = []
    for line in lines[header_index + 2 :]:
        if not line.startswith("|"):
            break
        rows.append(split_markdown_row(line))
    return rows


def field(
    name: str,
    field_type: str,
    mandatory: str,
    description: str,
    max_length: str = "N/A",
) -> tuple[str, str, str, str, str]:
    return name, field_type, mandatory, max_length, description


USE_CASE_META: dict[str, dict[str, object]] = {
    "UC-01": {
        "screens": "Login",
        "mockup": (
            "Draw the application identity, Email field, Password field with "
            "visibility toggle, Login button, loading state, inline validation, "
            "and authentication-error area."
        ),
        "fields": [
            field("Email", "Email text field", "Yes", "Firebase login email."),
            field(
                "Password",
                "Password field",
                "Yes",
                "Masked password with show or hide control.",
            ),
            field(
                "Login",
                "Button",
                "N/A",
                "Validates input, authenticates, and redirects by role.",
            ),
            field(
                "Authentication Error",
                "Message area",
                "N/A",
                "Displays invalid credentials, inactive account, or network errors.",
            ),
        ],
        "rules": "BR-01, BR-02, BR-23",
    },
    "UC-02": {
        "screens": "Admin Dashboard or User Profile",
        "mockup": (
            "Draw the Logout icon or button inside the Admin Dashboard app bar "
            "or User Profile screen and show the navigation result returning to Login."
        ),
        "fields": [
            field(
                "Logout",
                "Icon button or button",
                "N/A",
                "Ends the current Firebase Authentication session.",
            ),
            field(
                "Current User",
                "Read-only context",
                "Yes",
                "The authenticated account whose session will be ended.",
            ),
        ],
        "rules": "BR-01",
    },
    "UC-03": {
        "screens": "Change Password",
        "mockup": (
            "Draw Current Password, New Password, Confirm New Password, visibility "
            "controls, validation messages, Change Password button, and Back action. "
            "Do not draw Forgot Password, OTP, or password recovery."
        ),
        "fields": [
            field(
                "Current Password",
                "Password field",
                "Yes",
                "Used to re-authenticate the current user.",
            ),
            field(
                "New Password",
                "Password field",
                "Yes",
                "New Firebase Authentication password.",
            ),
            field(
                "Confirm New Password",
                "Password field",
                "Yes",
                "Must match New Password.",
            ),
            field(
                "Change Password",
                "Button",
                "N/A",
                "Re-authenticates and updates the password.",
            ),
        ],
        "rules": "The user must be authenticated and know the current password. Release 1.0 has no self-service Forgot Password screen.",
    },
    "UC-04": {
        "screens": "User Profile",
        "mockup": (
            "Draw avatar or identity header, read-only Email, editable Full Name, "
            "Phone, National ID, Date of Birth, Save Profile, theme control, and Logout."
        ),
        "fields": [
            field("Email", "Read-only email field", "Yes", "Login email; not editable."),
            field("Full Name", "Text field", "Yes", "Editable user full name."),
            field("Phone", "Phone field", "Yes", "Editable Vietnamese phone number."),
            field("National ID", "Text field", "Yes", "Editable identity number."),
            field(
                "Date of Birth",
                "Date picker",
                "No",
                "Editable date of birth.",
            ),
            field(
                "Save Profile",
                "Button",
                "N/A",
                "Validates and saves permitted profile fields.",
            ),
        ],
        "rules": "Role, status, apartment assignment, and email cannot be changed through self-service profile editing.",
    },
    "UC-05": {
        "screens": "Admin Home and Dashboard",
        "mockup": (
            "Draw the Admin Dashboard app bar, Refresh and Logout actions, five "
            "counter cards, quick-access cards, and Admin bottom navigation."
        ),
        "fields": [
            field("Apartment Count", "Read-only counter", "N/A", "Total apartments."),
            field("Resident Count", "Read-only counter", "N/A", "Total Resident profiles."),
            field(
                "Pending Requests",
                "Read-only counter",
                "N/A",
                "Requests with Pending status.",
            ),
            field("Unpaid Bills", "Read-only counter", "N/A", "Bills with Unpaid status."),
            field(
                "Visitors Inside",
                "Read-only counter",
                "N/A",
                "Visitors with Checked In status.",
            ),
            field("Refresh", "Icon button", "N/A", "Reloads all dashboard counters."),
        ],
        "rules": "Dashboard output is an on-screen summary and is not an exported report.",
    },
    "UC-06": {
        "screens": "User List, User Create, and User Edit",
        "mockup": (
            "Draw three frames: User List with search and filters; User Create with "
            "account fields; User Edit with role, apartment, and status controls."
        ),
        "fields": [
            field("Search", "Search field", "No", "Searches name, email, or apartment."),
            field("Role Filter", "Dropdown", "No", "Filters Admin, Staff, or Resident."),
            field("Status Filter", "Dropdown", "No", "Filters Active or Inactive."),
            field("Email", "Email field", "Yes", "Login email; immutable after creation."),
            field("Full Name", "Text field", "Yes", "User full name."),
            field("Phone", "Phone field", "Yes", "Vietnamese phone number."),
            field("National ID", "Text field", "Yes", "Identity number."),
            field("Role", "Dropdown", "Yes", "Admin, Staff, or Resident."),
            field("Apartment", "Dropdown", "No", "Optional apartment assignment."),
            field("Status", "Switch", "Yes", "Active or Inactive."),
            field("Create or Save", "Button", "N/A", "Creates or updates the account."),
        ],
        "rules": "BR-01, BR-02, BR-03, BR-04, BR-05",
    },
    "UC-07": {
        "screens": "Apartment List, Apartment Form, and Apartment Details",
        "mockup": (
            "Draw three frames: searchable Apartment List with floor and status filters; "
            "Apartment Form; Apartment Details with residents, Edit, Delete, and Assign Resident."
        ),
        "fields": [
            field("Search", "Search field", "No", "Searches apartment or owner name."),
            field("Floor Filter", "Dropdown", "No", "Filters by floor."),
            field("Status Filter", "Dropdown", "No", "Vacant or Occupied."),
            field("Apartment Number", "Text field", "Yes", "Apartment number."),
            field("Floor", "Number field", "Yes", "Floor number."),
            field("Building", "Text field", "Yes", "Building name."),
            field("Area", "Number field", "Yes", "Area in square metres."),
            field("Type", "Text field", "No", "Apartment layout."),
            field("Price", "Number field", "No", "Reference price."),
            field("Save", "Button", "N/A", "Creates or updates the apartment."),
            field("Delete", "Button", "N/A", "Deletes after confirmation."),
        ],
        "rules": "BR-06, BR-07, BR-22",
    },
    "UC-08": {
        "screens": "Apartment Details and Resident Picker Dialog",
        "mockup": (
            "Draw Apartment Details with the current resident list and an Assign Resident "
            "button. Draw the Resident Picker dialog used to select the Resident."
        ),
        "fields": [
            field(
                "Current Residents",
                "Read-only list",
                "N/A",
                "Resident identifiers assigned to the apartment.",
            ),
            field(
                "Owner Indicator",
                "Read-only label",
                "N/A",
                "Marks the Resident whose identifier equals ownerId.",
            ),
            field(
                "Assign Resident",
                "Button",
                "N/A",
                "Opens the Resident Picker dialog.",
            ),
            field(
                "Resident Selection",
                "Dialog list",
                "Yes",
                "Selects the Resident to assign.",
            ),
        ],
        "rules": "The current screen flow assigns the selected Resident and stores that Resident as apartment owner.",
    },
    "UC-09": {
        "screens": "Resident List, Resident Form, and Resident Profile",
        "mockup": (
            "Draw Resident List with search and status filter, Resident Form with identity "
            "and apartment fields, and read-only Resident Profile details."
        ),
        "fields": [
            field("Search", "Search field", "No", "Searches name, phone, or apartment."),
            field("Status Filter", "Dropdown", "No", "Active or Inactive."),
            field("Resident ID", "Text field", "Yes", "Firestore profile identifier."),
            field("Email", "Email field", "Yes", "Resident email."),
            field("Full Name", "Text field", "Yes", "Resident full name."),
            field("Phone", "Phone field", "Yes", "Resident phone."),
            field("National ID", "Text field", "Yes", "Resident identity number."),
            field("Apartment", "Dropdown", "No", "Apartment assignment."),
            field("Status", "Dropdown", "Yes", "Active or Inactive."),
            field("Save", "Button", "N/A", "Creates or updates the Resident profile."),
        ],
        "rules": "Creating a Resident profile in this feature currently writes Firestore profile data; account provisioning should be performed through User Management when login access is required.",
    },
    "UC-10": {
        "screens": "Request Create",
        "mockup": (
            "Draw Title, Category chips or dropdown, Description, image picker, image "
            "preview list with remove actions, image counter, and Submit Request."
        ),
        "fields": [
            field("Title", "Text field", "Yes", "Short request title."),
            field(
                "Category",
                "Choice control",
                "Yes",
                "Plumbing, Electrical, or General.",
            ),
            field("Description", "Multiline field", "Yes", "Issue description."),
            field("Images", "Image picker", "No", "Up to three images."),
            field("Submit Request", "Button", "N/A", "Uploads images and creates request."),
        ],
        "rules": "BR-08, BR-09. Resident must have an apartment assignment.",
    },
    "UC-11": {
        "screens": "My Request List and Request Details",
        "mockup": (
            "Draw a Resident request list with status badges and a detail frame containing "
            "category, description, images, status, processing account, and resolution note."
        ),
        "fields": [
            field("Request List", "Read-only list", "N/A", "Current Resident requests."),
            field("Status Badge", "Read-only badge", "N/A", "Current request status."),
            field("Request Details", "Read-only fields", "N/A", "Complete request information."),
            field("Resolution Note", "Read-only text", "No", "Staff or Admin resolution note."),
        ],
        "rules": "BR-08, BR-09",
    },
    "UC-12": {
        "screens": "Request Management and Request Details",
        "mockup": (
            "Draw status-filter tabs, request cards, request detail, status dropdown, "
            "resolution-note field, and Update Status button."
        ),
        "fields": [
            field("Status Filter", "Tabs or dropdown", "No", "Filters request status."),
            field("Request Details", "Read-only fields", "N/A", "Submitted request data."),
            field(
                "Status",
                "Dropdown",
                "Yes",
                "Pending, In Progress, or Completed.",
            ),
            field("Resolution Note", "Multiline field", "No", "Processing or completion note."),
            field("Update Status", "Button", "N/A", "Saves status and processing account."),
        ],
        "rules": "BR-09",
    },
    "UC-13": {
        "screens": "Complaint Create",
        "mockup": (
            "Draw a complaint or feedback multiline field, validation message, "
            "Resident and apartment context, and Submit Complaint button."
        ),
        "fields": [
            field("Content", "Multiline field", "Yes", "Complaint or feedback content."),
            field(
                "Resident and Apartment",
                "Read-only context",
                "Yes",
                "Derived from authenticated profile.",
            ),
            field("Submit Complaint", "Button", "N/A", "Creates a Submitted complaint."),
        ],
        "rules": "BR-10",
    },
    "UC-14": {
        "screens": "My Complaint List and Complaint Details",
        "mockup": (
            "Draw complaint cards with status badges and a details frame containing "
            "content, status, response, responder, and response time."
        ),
        "fields": [
            field("Complaint List", "Read-only list", "N/A", "Current Resident complaints."),
            field("Status Badge", "Read-only badge", "N/A", "Complaint status."),
            field("Response", "Read-only text", "No", "Admin or Staff response."),
            field("Response Time", "Read-only date-time", "No", "Time of response."),
        ],
        "rules": "BR-10",
    },
    "UC-15": {
        "screens": "Complaint Management and Complaint Details",
        "mockup": (
            "Draw status filter, complaint list, complaint details, Mark In Review, "
            "Response field, and Resolve action."
        ),
        "fields": [
            field("Status Filter", "Dropdown", "No", "Filters complaint status."),
            field("Complaint Details", "Read-only fields", "N/A", "Submitted complaint data."),
            field("Mark In Review", "Button", "N/A", "Changes status to In Review."),
            field("Response", "Multiline field", "Yes for resolve", "Response content."),
            field("Resolve", "Button", "N/A", "Stores response and resolves complaint."),
        ],
        "rules": "BR-10",
    },
    "UC-16": {
        "screens": "Bill Create",
        "mockup": (
            "Draw Apartment Identifier, Bill Type, Amount, Billing Month, Due Date, "
            "validation messages, and Create Bill."
        ),
        "fields": [
            field("Apartment Identifier", "Text field", "Yes", "Billed apartment."),
            field(
                "Bill Type",
                "Dropdown",
                "Yes",
                "Electricity, Water, Service, or Parking.",
            ),
            field("Amount", "Number field", "Yes", "Positive VND amount."),
            field("Billing Month", "Text field", "Yes", "YYYY-MM."),
            field("Due Date", "Date picker", "Yes", "Payment due date."),
            field("Create Bill", "Button", "N/A", "Creates an Unpaid bill."),
        ],
        "rules": "BR-11, BR-12. The current implementation must replace empty residentId and temporary createdBy values before production.",
    },
    "UC-17": {
        "screens": "Bill List and Bill Details",
        "mockup": (
            "Draw bill filters, bill cards, totals, and Bill Details containing bill data, "
            "payment information, and the actions available for the current status."
        ),
        "fields": [
            field("Apartment Filter", "Text or dropdown", "No", "Filters apartment."),
            field("Billing Month Filter", "Text or dropdown", "No", "Filters month."),
            field("Status Filter", "Dropdown", "No", "Filters bill status."),
            field("Bill List", "Read-only list", "N/A", "Matching bills."),
            field("Bill Details", "Read-only fields", "N/A", "Selected bill data."),
        ],
        "rules": "BR-11, BR-12",
    },
    "UC-18": {
        "screens": "My Bills",
        "mockup": (
            "Draw Resident apartment context, bill summary, status filter or grouping, "
            "bill cards with amount and due date, and navigation to Bill Payment."
        ),
        "fields": [
            field("Apartment", "Read-only text", "Yes", "Authenticated Resident apartment."),
            field("Bill List", "Read-only list", "N/A", "Bills for the apartment."),
            field("Bill Type", "Read-only label", "N/A", "Bill category."),
            field("Amount", "Read-only currency", "N/A", "Bill amount."),
            field("Due Date", "Read-only date", "N/A", "Bill due date."),
            field("Status", "Read-only badge", "N/A", "Bill status."),
        ],
        "rules": "BR-11, BR-12",
    },
    "UC-19": {
        "screens": "Bill Payment",
        "mockup": (
            "Draw bill summary, bank name, account number, transfer reference, copy actions, "
            "payment amount, confirmation explanation, and Confirm Transfer button."
        ),
        "fields": [
            field("Bill Summary", "Read-only fields", "Yes", "Selected bill information."),
            field("Bank Name", "Read-only text", "Yes", "Building bank."),
            field("Account Number", "Read-only text", "Yes", "Destination account."),
            field("Transfer Reference", "Read-only text", "Yes", "Required transfer content."),
            field("Copy", "Icon buttons", "N/A", "Copies bank information."),
            field("Confirm Transfer", "Button", "N/A", "Creates a Pending payment request."),
        ],
        "rules": "BR-13, BR-14, BR-15. The current proof reference is mock data and no complete receipt-upload flow exists.",
    },
    "UC-20": {
        "screens": "Bill Details — Record Cash Payment",
        "mockup": (
            "Draw the Bill Details state for an Unpaid bill, Record Cash Payment button, "
            "confirmation dialog, and successful Paid state."
        ),
        "fields": [
            field("Bill Details", "Read-only fields", "Yes", "Selected unpaid bill."),
            field(
                "Record Cash Payment",
                "Button",
                "N/A",
                "Opens cash-receipt confirmation.",
            ),
            field("Confirmation", "Dialog", "Yes", "Confirms cash was received."),
        ],
        "rules": "BR-13, BR-14, BR-18. The current temporary Staff identifier must be replaced before production.",
    },
    "UC-21": {
        "screens": "Bill Details — Pending Bank Transfer",
        "mockup": (
            "Draw Pending payment information, amount, payment method, proof reference, "
            "Approve button, Reject button, and rejection-reason dialog."
        ),
        "fields": [
            field("Pending Payment", "Read-only fields", "Yes", "Payment awaiting verification."),
            field("Proof Reference", "Read-only text", "No", "Stored transfer proof reference."),
            field("Approve", "Button", "N/A", "Approves payment and marks bill Paid."),
            field("Reject", "Button", "N/A", "Opens rejection-reason dialog."),
            field("Rejection Reason", "Multiline field", "Yes for reject", "Required reason."),
        ],
        "rules": "BR-14, BR-16, BR-17. Temporary Staff identifiers and mock proof data must be replaced before production.",
    },
    "UC-22": {
        "screens": "Payment History",
        "mockup": (
            "Draw a chronological payment list with amount, method, status, date-time, "
            "and empty or error state."
        ),
        "fields": [
            field("Payment List", "Read-only list", "N/A", "Resident or apartment payments."),
            field("Amount", "Read-only currency", "N/A", "Payment amount."),
            field("Method", "Read-only label", "N/A", "Cash or Bank Transfer."),
            field("Status", "Read-only badge", "N/A", "Pending, Approved, or Rejected."),
            field("Date and Time", "Read-only date-time", "N/A", "Payment creation time."),
        ],
        "rules": "BR-13, BR-14, BR-23",
    },
    "UC-23": {
        "screens": "Announcement List and Announcement Details",
        "mockup": (
            "Draw announcement cards with title, type, and date, followed by an Announcement "
            "Details frame. Show create or edit controls only for Admin and Staff."
        ),
        "fields": [
            field("Announcement List", "Read-only list", "N/A", "In-app announcements."),
            field("Title", "Read-only text", "Yes", "Announcement title."),
            field("Content", "Read-only text", "Yes", "Announcement content."),
            field("Type", "Read-only label", "Yes", "Announcement type."),
            field("Created Time", "Read-only date-time", "No", "Creation time."),
        ],
        "rules": "Release 1.0 does not filter retrieval by targetRoles and does not send FCM push notifications.",
    },
    "UC-24": {
        "screens": "Announcement Create or Edit and Announcement Details",
        "mockup": (
            "Draw Title, Content, Type, Target Roles, Save, Edit, Delete, and the "
            "delete-confirmation dialog."
        ),
        "fields": [
            field("Title", "Text field", "Yes", "Announcement title."),
            field("Content", "Multiline field", "Yes", "Announcement content."),
            field("Type", "Text or dropdown", "Yes", "Announcement type."),
            field("Target Roles", "Multi-select", "Yes", "Stored role metadata."),
            field("Save", "Button", "N/A", "Creates or updates announcement."),
            field("Delete", "Button", "N/A", "Deletes after confirmation."),
        ],
        "rules": "BR-20, BR-22",
    },
    "UC-25": {
        "screens": "Visitor Registration",
        "mockup": (
            "Draw Visitor Name, Phone, Purpose, Expected Time, apartment context, "
            "validation messages, and Register Visitor."
        ),
        "fields": [
            field("Visitor Name", "Text field", "Yes", "Visitor full name."),
            field("Visitor Phone", "Phone field", "Yes", "Visitor phone number."),
            field("Purpose", "Text field", "Yes", "Purpose of visit."),
            field("Expected Time", "Date-time picker", "No", "Expected arrival."),
            field("Apartment", "Read-only or selected value", "Yes", "Destination apartment."),
            field("Register Visitor", "Button", "N/A", "Creates Registered visitor."),
        ],
        "rules": "BR-21",
    },
    "UC-26": {
        "screens": "Visitor List",
        "mockup": (
            "Draw search, visitor cards, identity, apartment, expected time, status badge, "
            "check-in time, check-out time, and context-sensitive actions."
        ),
        "fields": [
            field("Search", "Search field", "No", "Searches visitor information."),
            field("Visitor List", "Read-only list", "N/A", "Registered visitors."),
            field("Apartment", "Read-only text", "Yes", "Destination apartment."),
            field("Expected Time", "Read-only date-time", "No", "Expected arrival."),
            field("Status", "Read-only badge", "Yes", "Visitor status."),
        ],
        "rules": "BR-21",
    },
    "UC-27": {
        "screens": "Visitor List — Check-In and Check-Out Actions",
        "mockup": (
            "Draw Registered state with Check In, Checked In state with Check Out, "
            "Checked Out final state, timestamps, and operation feedback."
        ),
        "fields": [
            field("Visitor Details", "Read-only fields", "Yes", "Selected visitor."),
            field("Check In", "Button", "N/A", "Available only for Registered."),
            field("Check Out", "Button", "N/A", "Available only for Checked In."),
            field("Check-In Time", "Read-only date-time", "No", "Recorded arrival."),
            field("Check-Out Time", "Read-only date-time", "No", "Recorded departure."),
        ],
        "rules": "BR-21",
    },
}


FEATURES: list[tuple[str, str, list[str]]] = [
    ("3.2", "Authentication and Profile Management", ["UC-01", "UC-02", "UC-03", "UC-04"]),
    ("3.3", "Dashboard", ["UC-05"]),
    ("3.4", "User Management", ["UC-06"]),
    ("3.5", "Apartment and Resident Management", ["UC-07", "UC-08", "UC-09"]),
    ("3.6", "Maintenance Request Management", ["UC-10", "UC-11", "UC-12"]),
    ("3.7", "Complaint and Feedback Management", ["UC-13", "UC-14", "UC-15"]),
    (
        "3.8",
        "Bill and Manual Payment Management",
        ["UC-16", "UC-17", "UC-18", "UC-19", "UC-20", "UC-21", "UC-22"],
    ),
    ("3.9", "Announcement Management", ["UC-23", "UC-24"]),
    ("3.10", "Visitor Management", ["UC-25", "UC-26", "UC-27"]),
]


def markdown_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    output = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join("---" for _ in headers) + "|",
    ]
    for row in rows:
        escaped = [value.replace("|", r"\|").replace("\n", "<br>") for value in row]
        output.append("| " + " | ".join(escaped) + " |")
    return output


def build_template_markdown() -> str:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    feature_start = source_lines.index("### 3.2 Authentication and Profile Management")
    nfr_start = source_lines.index("## 4. Non-Functional Requirements")
    prefix = source_lines[:feature_start]
    suffix = source_lines[nfr_start:]

    use_case_rows = read_table_after(
        source_lines,
        "| ID | Function Group | Use Case | Actors |",
    )
    condition_rows = read_table_after(
        source_lines,
        "| Use Case | Preconditions | Trigger | Postconditions |",
    )
    use_cases = {
        row[0]: {
            "group": row[1],
            "name": row[2],
            "actors": row[3],
            "description": row[4],
        }
        for row in use_case_rows
    }
    conditions = {
        row[0].split()[0]: {
            "preconditions": row[1],
            "trigger": row[2],
            "postconditions": row[3],
            "alternatives": row[4],
        }
        for row in condition_rows
    }

    output = list(prefix)
    output.extend(
        [
            "",
            "The feature details below follow the required SRS template. Each function contains a reserved screen mock-up area, a screen definition table, and a detailed use case description.",
            "",
        ]
    )

    figure_index = 1
    for feature_number, feature_name, uc_ids in FEATURES:
        output.extend([f"### {feature_number} {feature_name}", ""])
        for function_index, uc_id in enumerate(uc_ids, start=1):
            uc = use_cases[uc_id]
            condition = conditions[uc_id]
            meta = USE_CASE_META[uc_id]
            function_number = f"{feature_number}.{function_index}"
            output.extend(
                [
                    f"#### {function_number} {uc_id} — {uc['name']}",
                    "",
                    f"##### {function_number}.1 Screen Mock-up",
                    "",
                    f"[[FIGURE_MOCKUP|3-{figure_index}|{meta['screens']}]]",
                    "",
                    f"**Table {function_number}-1: Screen Definition**",
                    "",
                ]
            )
            figure_index += 1
            field_rows = [
                [
                    str(index),
                    values[0],
                    values[1],
                    values[2],
                    values[3],
                    values[4],
                ]
                for index, values in enumerate(meta["fields"], start=1)
            ]
            output.extend(
                markdown_table(
                    [
                        "#",
                        "Field Name",
                        "Type",
                        "Mandatory",
                        "Max Length",
                        "Description",
                    ],
                    field_rows,
                )
            )
            output.extend(
                [
                    "",
                    f"##### {function_number}.2 Use Case Description",
                    "",
                    f"[[USE_CASE_DETAIL|{uc_id}|{feature_name}]]",
                    "",
                ]
            )

    output.extend(suffix)
    return "\n".join(output).rstrip() + "\n"


def add_inline(paragraph, text: str, size: float | None = None) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            if size is not None:
                run.font.size = Pt(size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.color.rgb = RGBColor(80, 80, 80)
        if size is not None:
            run.font.size = Pt(size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if size is not None:
            run.font.size = Pt(size)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 70, start: int = 80, bottom: int = 70, end: int = 80
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))


def column_widths(column_count: int) -> list[float]:
    presets = {
        2: [1.35, 5.75],
        3: [0.50, 1.90, 4.70],
        4: [0.50, 1.45, 1.70, 3.45],
        5: [0.48, 1.12, 1.30, 1.00, 3.20],
        6: [0.35, 1.20, 0.75, 0.70, 0.70, 3.40],
    }
    return presets.get(column_count, [7.1 / column_count] * column_count)


def is_separator_row(cells: list[str]) -> bool:
    return all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells
    )


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = column_widths(column_count)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            repeat_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "FCE4D6")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif column_index == 0 and column_count >= 3:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value.replace("<br>", "\n"), size=8.25)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                if row_index == 0:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_mockup_area(document: Document, uc_id: str, screen_name: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.1)
    row = table.rows[0]
    row.height = Cm(14.0 if uc_id.startswith("E.") else 7.2)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    prevent_row_split(row)
    cell = row.cells[0]
    cell.width = Inches(7.1)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F8FAFC")
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("RESERVED SCREEN MOCK-UP AREA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.add_run(f"\n{uc_id} — {screen_name}")
    instruction = paragraph.add_run(
        "\n\nPaste or insert the hand-drawn screen mock-up inside this bordered area."
    )
    instruction.italic = True
    instruction.font.size = Pt(9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure_mockup_area(
    document: Document,
    figure_number: str,
    screen_name: str,
) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(5)
    run = caption.add_run(f"Figure {figure_number}: Screen Design of {screen_name}")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.5)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.1)
    row = table.rows[0]
    row.height = Cm(8.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    prevent_row_split(row)
    cell = row.cells[0]
    cell.width = Inches(7.1)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paste_run = paragraph.add_run("DÁN ẢNH Ở ĐÂY")
    paste_run.bold = True
    paste_run.font.name = "Arial"
    paste_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    paste_run.font.size = Pt(14)
    paste_run.font.color.rgb = RGBColor(127, 127, 127)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def merge_range(row, start: int, end: int):
    cell = row.cells[start]
    for index in range(start + 1, end + 1):
        cell = cell.merge(row.cells[index])
    return cell


def set_table_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    if fill is not None:
        shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if align is not None:
        paragraph.alignment = align
    add_inline(paragraph, text, size=9.0)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if bold:
            run.bold = True


def add_template_table(document: Document, rows: int, columns: int):
    table = document.add_table(rows=rows, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        prevent_row_split(row)
    return table


def split_sentences(text: str) -> list[str]:
    sentences = [
        item.strip()
        for item in re.split(r"(?<=[.!?])\s+(?=[A-Z])", text.strip())
        if item.strip()
    ]
    return sentences or [text.strip()]


def flow_actor(action: str, actors: str) -> str:
    lowered = action.lower()
    if lowered.startswith("the system") or lowered.startswith("system"):
        return "System"
    if lowered.startswith("admin or staff"):
        return "Admin/Staff"
    if lowered.startswith("admin"):
        return "Admin"
    if lowered.startswith("staff"):
        return "Staff"
    if lowered.startswith("the resident") or lowered.startswith("resident"):
        return "Resident"
    if lowered.startswith("the authenticated user") or lowered.startswith("the user"):
        return "User"
    return actors.replace(", ", "/")


def load_detailed_use_case_data():
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    use_case_rows = read_table_after(
        source_lines,
        "| ID | Function Group | Use Case | Actors |",
    )
    condition_rows = read_table_after(
        source_lines,
        "| Use Case | Preconditions | Trigger | Postconditions |",
    )
    rule_rows = read_table_after(source_lines, "| ID | Rule Definition |")
    use_cases = {
        row[0]: {
            "name": row[2],
            "actors": row[3],
            "description": row[4],
        }
        for row in use_case_rows
    }
    conditions = {
        row[0].split()[0]: {
            "preconditions": row[1],
            "trigger": row[2],
            "postconditions": row[3],
            "alternatives": row[4],
        }
        for row in condition_rows
    }
    rules = {row[0]: row[1] for row in rule_rows}
    return use_cases, conditions, rules


def add_labeled_detail_row(table, label: str, value: str) -> None:
    row = table.add_row()
    merge_range(row, 0, 1)
    merge_range(row, 2, 6)
    set_table_cell(row.cells[0], label, bold=True, fill="FCE4D6")
    set_table_cell(row.cells[2], value)


def add_use_case_description(
    document: Document,
    uc_id: str,
    feature_name: str,
) -> None:
    use_cases, conditions, rule_definitions = load_detailed_use_case_data()
    uc = use_cases[uc_id]
    condition = conditions[uc_id]

    header = add_template_table(document, 2, 7)
    merge_range(header.rows[0], 0, 1)
    merge_range(header.rows[0], 3, 4)
    merge_range(header.rows[0], 5, 6)
    merge_range(header.rows[1], 0, 1)
    header_values = [
        (header.rows[0].cells[0], "Use Case ID", True, "FCE4D6"),
        (header.rows[0].cells[2], uc_id, True, None),
        (header.rows[0].cells[3], "Use Case Name", True, "FCE4D6"),
        (header.rows[0].cells[5], uc["name"], True, None),
        (header.rows[1].cells[0], "Author", True, "FCE4D6"),
        (header.rows[1].cells[2], "Group 5", True, None),
        (header.rows[1].cells[3], "Version", True, "FCE4D6"),
        (header.rows[1].cells[4], "1.0", True, None),
        (header.rows[1].cells[5], "Date", True, "FCE4D6"),
        (header.rows[1].cells[6], "23/07/2026", True, None),
    ]
    for cell, value, bold, fill in header_values:
        set_table_cell(cell, value, bold=bold, fill=fill)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    detail = add_template_table(document, 0, 7)
    add_labeled_detail_row(detail, "Actor", uc["actors"])
    add_labeled_detail_row(detail, "Description", uc["description"])
    add_labeled_detail_row(
        detail,
        "Precondition",
        "\n".join(
            f"PRE-{index:02d}: {sentence}"
            for index, sentence in enumerate(
                split_sentences(condition["preconditions"]),
                start=1,
            )
        ),
    )
    add_labeled_detail_row(detail, "Trigger", f"TRG-01: {condition['trigger']}")
    add_labeled_detail_row(
        detail,
        "Post-Condition",
        "\n".join(
            f"POS-{index:02d}: {sentence}"
            for index, sentence in enumerate(
                split_sentences(condition["postconditions"]),
                start=1,
            )
        ),
    )

    main_heading = detail.add_row()
    merge_range(main_heading, 0, 6)
    set_table_cell(
        main_heading.cells[0],
        "Main flows",
        bold=True,
        fill="FCE4D6",
    )
    main_header = detail.add_row()
    merge_range(main_header, 2, 6)
    set_table_cell(main_header.cells[0], "Step", bold=True, fill="FCE4D6")
    set_table_cell(main_header.cells[1], "Actor", bold=True, fill="FCE4D6")
    set_table_cell(main_header.cells[2], "Action", bold=True, fill="FCE4D6")
    for step, action in enumerate(split_sentences(uc["description"]), start=1):
        row = detail.add_row()
        merge_range(row, 2, 6)
        set_table_cell(row.cells[0], str(step), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_cell(
            row.cells[1],
            flow_actor(action, uc["actors"]),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_table_cell(row.cells[2], action)

    alternative_heading = document.add_paragraph()
    alternative_run = alternative_heading.add_run("Alternative flows")
    alternative_run.bold = True
    alternative_run.font.name = "Arial"
    alternative_run.font.size = Pt(10.5)
    alternative_heading.paragraph_format.keep_with_next = True

    alternative = add_template_table(document, 3, 3)
    merge_range(alternative.rows[0], 1, 2)
    set_table_cell(
        alternative.rows[0].cells[0],
        "AT1",
        bold=True,
        fill="FCE4D6",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_table_cell(
        alternative.rows[0].cells[1],
        condition["alternatives"],
    )
    set_table_cell(alternative.rows[1].cells[0], "Sub step", bold=True, fill="FCE4D6")
    set_table_cell(alternative.rows[1].cells[1], "Actor", bold=True, fill="FCE4D6")
    set_table_cell(alternative.rows[1].cells[2], "Action", bold=True, fill="FCE4D6")
    set_table_cell(
        alternative.rows[2].cells[0],
        "1.1",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_table_cell(
        alternative.rows[2].cells[1],
        "System",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_table_cell(
        alternative.rows[2].cells[2],
        "Display an appropriate error or empty state, do not complete the requested "
        "operation, and keep the current stored data unchanged.",
    )

    business_heading = document.add_paragraph()
    business_run = business_heading.add_run("Business Rules")
    business_run.bold = True
    business_run.font.name = "Arial"
    business_run.font.size = Pt(10.5)
    business_heading.paragraph_format.keep_with_next = True

    business = add_template_table(document, 1, 2)
    set_table_cell(
        business.rows[0].cells[0],
        "#",
        bold=True,
        fill="FCE4D6",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_table_cell(
        business.rows[0].cells[1],
        "Rule Description",
        bold=True,
        fill="FCE4D6",
    )
    rule_reference = str(USE_CASE_META[uc_id]["rules"])
    rule_ids = re.findall(r"BR-\d{2}", rule_reference)
    if rule_ids:
        for rule_id in rule_ids:
            row = business.add_row()
            set_table_cell(
                row.cells[0],
                rule_id,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_table_cell(
                row.cells[1],
                rule_definitions.get(rule_id, "See Section 5.1 Business Rules."),
            )
    else:
        row = business.add_row()
        set_table_cell(
            row.cells[0],
            f"BR-{uc_id}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_table_cell(row.cells[1], rule_reference)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_docx(markdown: str) -> None:
    lines = markdown.splitlines()
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for list_style_name in ("List Bullet", "List Number"):
        list_style = document.styles[list_style_name]
        list_style.font.name = "Times New Roman"
        list_style.font.size = Pt(11)
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    heading_specs = {
        "Title": (16, RGBColor(31, 78, 121)),
        "Heading 1": (15, RGBColor(31, 78, 121)),
        "Heading 2": (13, RGBColor(31, 78, 121)),
        "Heading 3": (12, RGBColor(55, 55, 55)),
        "Heading 4": (11, RGBColor(55, 55, 55)),
        "Heading 5": (10.5, RGBColor(55, 55, 55)),
    }
    for style_name, (size, color) in heading_specs.items():
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    document.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    index = 0
    first_heading = True
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or stripped == "---":
            index += 1
            continue

        mockup_match = re.fullmatch(
            r"\[\[MOCKUP_AREA\|([^|]+)\|(.+)\]\]", stripped
        )
        if mockup_match:
            add_mockup_area(document, mockup_match.group(1), mockup_match.group(2))
            index += 1
            continue

        figure_mockup_match = re.fullmatch(
            r"\[\[FIGURE_MOCKUP\|([^|]+)\|(.+)\]\]",
            stripped,
        )
        if figure_mockup_match:
            add_figure_mockup_area(
                document,
                figure_mockup_match.group(1),
                figure_mockup_match.group(2),
            )
            index += 1
            continue

        use_case_match = re.fullmatch(
            r"\[\[USE_CASE_DETAIL\|([^|]+)\|(.+)\]\]",
            stripped,
        )
        if use_case_match:
            add_use_case_description(
                document,
                use_case_match.group(1),
                use_case_match.group(2),
            )
            index += 1
            continue

        if stripped.startswith("|"):
            block: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = split_markdown_row(lines[index])
                if not is_separator_row(cells):
                    block.append(cells)
                index += 1
            add_table(document, block)
            continue

        heading_match = re.match(r"^(#{1,5})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if (
                text in {
                    "II. Software Requirement Specification",
                    "Diagram Redrawing Guide",
                    "Final Consistency Checklist",
                }
                and not first_heading
            ):
                document.add_page_break()
            if re.match(r"^3\.(?:[2-9]|10)\s", text):
                document.add_page_break()
            if re.match(r"^E\.\d+\s", text):
                document.add_page_break()
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                first_heading = False
            else:
                paragraph = document.add_paragraph(
                    style=f"Heading {min(level, 5)}"
                )
            add_inline(paragraph, text)
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet_match.group(1))
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, number_match.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1

    for document_section in document.sections:
        footer_paragraph = document_section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.add_run("Page ")
        field_node = OxmlElement("w:fldSimple")
        field_node.set(qn("w:instr"), "PAGE")
        footer_paragraph._p.append(field_node)

    document.core_properties.title = (
        "Template-Aligned Section II — Software Requirement Specification"
    )
    document.core_properties.subject = "Apartment Building Management System"
    document.core_properties.author = "Group 5"
    document.core_properties.comments = (
        "Contains 27 hand-drawn screen insertion frames and template-aligned "
        "detailed use case descriptions."
    )
    document.save(DOCX_TARGET)


def validate_outputs() -> None:
    markdown = MARKDOWN_TARGET.read_text(encoding="utf-8")
    if markdown.count("[[FIGURE_MOCKUP|") != 27:
        raise RuntimeError("The Markdown output does not contain 27 figure mock-up areas.")
    for uc_number in range(1, 28):
        uc_id = f"UC-{uc_number:02d}"
        if uc_id not in markdown:
            raise RuntimeError(f"Missing {uc_id} from Markdown output.")

    with ZipFile(DOCX_TARGET) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"Corrupt DOCX member: {bad_member}")

    document = Document(DOCX_TARGET)
    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    if all_text.count("DÁN ẢNH Ở ĐÂY") != 27:
        raise RuntimeError("The DOCX output does not contain 27 image insertion boxes.")
    if "Screen or UI scope:" in all_text:
        raise RuntimeError("The DOCX output still contains Screen or UI scope text.")
    if "Drawing requirements:" in all_text:
        raise RuntimeError("The DOCX output still contains Drawing requirements text.")
    if "Mock-up reference:" in all_text:
        raise RuntimeError("The DOCX output still contains Mock-up reference text.")
    if sum(
        1
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Alternative flows"
    ) != 27:
        raise RuntimeError("The DOCX output does not contain 27 Alternative flows sections.")
    if sum(
        1
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Business Rules"
    ) != 27:
        raise RuntimeError("The DOCX output does not contain 27 Business Rules sections.")
    for uc_number in range(1, 28):
        uc_id = f"UC-{uc_number:02d}"
        if uc_id not in all_text:
            raise RuntimeError(f"Missing {uc_id} from DOCX output.")

    print(f"Markdown: {MARKDOWN_TARGET}")
    print(f"DOCX: {DOCX_TARGET}")
    print(f"DOCX size: {DOCX_TARGET.stat().st_size} bytes")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")
    print("Mock-up areas: 27")
    print("Use cases: 27")


def main() -> None:
    markdown = build_template_markdown()
    MARKDOWN_TARGET.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    validate_outputs()


if __name__ == "__main__":
    main()
